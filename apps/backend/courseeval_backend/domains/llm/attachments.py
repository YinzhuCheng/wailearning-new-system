from __future__ import annotations

import base64
import io
import json
import mimetypes
import os
import re
import shutil
import subprocess
import tempfile
import uuid
import zipfile
from dataclasses import dataclass
from pathlib import Path, PurePosixPath
from typing import Any, Optional

import fitz
import olefile
import openpyxl
import rarfile
import xlrd
from docx import Document
from PIL import Image, ImageFile, UnidentifiedImageError

from apps.backend.courseeval_backend.attachments import get_attachment_file_path
from apps.backend.courseeval_backend.core.config import settings

def _rar_extractor_tool_path() -> Optional[tuple[str, str]]:
    unrar_tool = shutil.which("unrar") or shutil.which("unrar-free")
    if unrar_tool:
        return "unrar", unrar_tool
    bsdtar_tool = shutil.which("bsdtar")
    if bsdtar_tool:
        return "tar", bsdtar_tool
    tar_tool = shutil.which("tar")
    if tar_tool:
        try:
            proc = subprocess.run(
                [tar_tool, "--version"],
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                timeout=5,
            )
        except (OSError, subprocess.TimeoutExpired):
            proc = None
        output = (proc.stdout if proc else b"").decode("utf-8", errors="replace").lower()
        if "libarchive" in output or "bsdtar" in output:
            return "tar", tar_tool
    return None


def _unrar_tool_path() -> Optional[str]:
    tool = _rar_extractor_tool_path()
    return tool[1] if tool else None


def _rar_read_member_bytes(archive_path: str, member_name: str) -> bytes:
    """Extract one RAR member via unrar/unrar-free or libarchive-backed tar."""
    import shutil

    tool_info = _rar_extractor_tool_path()
    tool = tool_info[1] if tool_info else None
    if not tool:
        raise RuntimeError("未找到 unrar / unrar-free / tar，无法解压 RAR。")
    abs_arc = os.path.abspath(archive_path)
    norm_member = (member_name or "").replace("\\", "/")
    if tool_info[0] == "tar":
        tmp_path_obj = Path(tempfile.gettempdir()) / f"rar-one-{uuid.uuid4().hex}"
        tmp_path_obj.mkdir(parents=True, exist_ok=False)
        tmp_dir = str(tmp_path_obj)
    else:
        tmp_dir = tempfile.mkdtemp(prefix="rar-one-")
    try:
        if tool_info[0] == "tar":
            command = [tool, "-xf", abs_arc, "-C", tmp_dir, norm_member]
            cwd = None
        else:
            command = [tool, "x", "-o+", "-y", abs_arc, norm_member]
            cwd = tmp_dir
        proc = subprocess.run(
            command,
            cwd=cwd,
            capture_output=True,
            timeout=120,
        )
        if proc.returncode != 0:
            err = (proc.stderr or proc.stdout or b"").decode("utf-8", errors="replace")[:500]
            raise RuntimeError(f"unrar 解压失败（{proc.returncode}）：{err}")
        out_path = os.path.join(tmp_dir, norm_member)
        if not os.path.isfile(out_path):
            base = os.path.basename(norm_member)
            alt = os.path.join(tmp_dir, base)
            out_path = alt if os.path.isfile(alt) else out_path
        if not os.path.isfile(out_path):
            raise RuntimeError("unrar 解压后未找到目标文件。")
        return Path(out_path).read_bytes()
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

VISION_TEST_IMAGE_DATA_URL = (
    "data:image/png;base64,"
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8"
    "/w8AAusB9Y9nKXUAAAAASUVORK5CYII="
)
# Cap encoded image size for multi-modal test requests (avoids huge payloads to LLM APIs).
MAX_VISION_TEST_IMAGE_BYTES = 5 * 1024 * 1024


def build_png_data_url_from_image_bytes(data: bytes) -> str:
    """Load common formats (jpeg/png/webp/gif/bmp) and emit an OpenAI-compatible data:image/png;base64,... URL."""
    if not data or len(data) > MAX_VISION_TEST_IMAGE_BYTES:
        raise ValueError(f"Image must be non-empty and at most {MAX_VISION_TEST_IMAGE_BYTES} bytes.")
    prev_truncated = ImageFile.LOAD_TRUNCATED_IMAGES
    ImageFile.LOAD_TRUNCATED_IMAGES = True
    try:
        im = Image.open(io.BytesIO(data))
        im.load()
    except (UnidentifiedImageError, OSError) as exc:
        raise ValueError("无法将文件识别为支持的图片（请使用 JPEG/PNG/WebP 等）。") from exc
    finally:
        ImageFile.LOAD_TRUNCATED_IMAGES = prev_truncated
    if im.mode not in ("RGB", "RGBA"):
        im = im.convert("RGBA" if (getattr(im, "info", None) and im.info.get("transparency") is not None) else "RGB")
    out = io.BytesIO()
    im.save(out, format="PNG", optimize=True)
    raw = out.getvalue()
    if len(raw) > MAX_VISION_TEST_IMAGE_BYTES:
        raise ValueError("转码为 PNG 后仍过大，请使用更小的图片。")
    b64 = base64.b64encode(raw).decode("ascii")
    return f"data:image/png;base64,{b64}"


JSON_FENCE_PATTERN = re.compile(r"^\s*```(?:json)?\s*(.*?)\s*```\s*$", re.DOTALL | re.IGNORECASE)
MAX_ZIP_DEPTH = 4
MAX_ZIP_FILES = 100
MAX_ZIP_TOTAL_BYTES = 80 * 1024 * 1024
MAX_FILE_TEXT_CHARS = 12000
MAX_IPYNB_OUTPUT_CHARS = 6000
# Prior attempts included in the grading prompt (text-only summary); older rounds omitted for token cost.
ITERATION_CONTEXT_MAX_PRIOR_ATTEMPTS = 2
ITERATION_PRIOR_NOTE_CHAR_MAX = 900
ITERATION_PRIOR_COMMENT_CHAR_MAX = 500
SUPPORTED_TEXT_EXTENSIONS = {
    ".c",
    ".cc",
    ".cpp",
    ".csv",
    ".docx",
    ".go",
    ".h",
    ".hpp",
    ".html",
    ".java",
    ".js",
    ".json",
    ".jsx",
    ".md",
    ".py",
    ".rb",
    ".rs",
    ".sql",
    ".tex",
    ".ts",
    ".tsx",
    ".txt",
    ".vue",
    ".xml",
    ".yaml",
    ".yml",
}
SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}
ZIP_EXTENSIONS = {".zip"}
RAR_EXTENSIONS = {".rar"}
EXCEL_XLSX_EXTENSIONS = {".xlsx"}
EXCEL_XLS_EXTENSIONS = {".xls"}
LEGACY_WORD_DOC_EXTENSIONS = {".doc"}
IPYNB_EXTENSIONS = {".ipynb"}
PDF_EXTENSIONS = {".pdf"}

@dataclass
class MaterialBlock:
    priority: int
    path: str
    block_type: str
    text: Optional[str] = None
    image_data_url: Optional[str] = None
    estimated_tokens: int = 0
    logical_path: Optional[str] = None
    mime_hint: Optional[str] = None
    origin: Optional[str] = None
    truncated: bool = False

def _truncate_text(value: str, limit: int = MAX_FILE_TEXT_CHARS) -> tuple[str, bool]:
    if len(value) <= limit:
        return value, False
    return value[:limit], True


def _decode_bytes_as_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _safe_relative_path(path_text: str) -> Optional[str]:
    normalized = PurePosixPath(path_text.replace("\\", "/"))
    safe_parts = []
    for part in normalized.parts:
        if part in {"", ".", ".."}:
            continue
        safe_parts.append(part)
    if not safe_parts:
        return None
    return "/".join(safe_parts)


def _guess_mime_type(path: str) -> str:
    mime_type, _ = mimetypes.guess_type(path)
    return mime_type or "application/octet-stream"


def _bytes_to_data_url(path: str, content: bytes) -> str:
    mime_type = _guess_mime_type(path)
    return f"data:{mime_type};base64,{base64.b64encode(content).decode('ascii')}"


def _extract_docx_text(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        t = (paragraph.text or "").strip()
        if t:
            parts.append(t)
    for table in document.tables:
        rows_out: list[str] = []
        for row in table.rows:
            cells = [" ".join((c.text or "").split()) for c in row.cells]
            if any(cells):
                rows_out.append("\t".join(cells))
        if rows_out:
            parts.append("[表格]\n" + "\n".join(rows_out))
    return "\n".join(parts)


def _extract_xlsx_text(content: bytes, path: str) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    try:
        parts: list[str] = []
        max_rows = 300
        max_cols = 64
        for sheet in wb.worksheets:
            parts.append(f"[工作表: {sheet.title}]")
            for row in sheet.iter_rows(min_row=1, max_row=max_rows, max_col=max_cols, values_only=True):
                cells = ["" if v is None else str(v).strip() for v in row]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts).strip()
    finally:
        wb.close()


def _extract_xls_text(content: bytes, path: str) -> str:
    book = xlrd.open_workbook(file_contents=content, formatting_info=False, on_demand=True)
    parts: list[str] = []
    max_rows = 300
    max_cols = 64
    try:
        for sheet in book.sheets():
            parts.append(f"[工作表: {sheet.name}]")
            for rx in range(min(sheet.nrows, max_rows)):
                row_vals = []
                for cx in range(min(sheet.ncols, max_cols)):
                    row_vals.append(str(sheet.cell_value(rx, cx)).strip())
                if any(row_vals):
                    parts.append("\t".join(row_vals))
    finally:
        book.release_resources()
    return "\n".join(parts).strip()


def _extract_legacy_doc_text(content: bytes) -> str:
    """
    Best-effort text from legacy .doc (OLE compound file). Does not preserve layout/tables.
    Reads the main Word binary streams only (avoids pulling unrelated OLE blobs).
    """
    if not olefile.isOleFile(io.BytesIO(content)):
        return ""
    ole = olefile.OleFileIO(io.BytesIO(content))
    max_stream_bytes = 6 * 1024 * 1024
    blob = b""
    try:
        for s in ole.listdir():
            joined = "/".join(s).lower()
            if not (
                "worddocument" in joined
                or joined.endswith("1table")
                or joined.endswith("0table")
                or joined == "data"
            ):
                continue
            try:
                blob += ole.openstream(s).read(max_stream_bytes)
            except Exception:
                continue
        if not blob:
            for s in ole.listdir():
                try:
                    blob += ole.openstream(s).read(min(max_stream_bytes, 512 * 1024))
                    if len(blob) >= max_stream_bytes:
                        break
                except Exception:
                    continue
    finally:
        ole.close()
    if not blob:
        return ""
    text = blob.decode("utf-16le", errors="ignore")
    lines: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if len(line) < 2:
            continue
        printable = sum(1 for c in line if c.isprintable() or c in "\t")
        if printable / max(len(line), 1) < 0.55:
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def _extract_pdf_images(content: bytes, path: str) -> list[MaterialBlock]:
    blocks: list[MaterialBlock] = []
    document = fitz.open(stream=content, filetype="pdf")
    try:
        for index, page in enumerate(document, start=1):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            image_bytes = pixmap.tobytes("png")
            logical = f"{path}#page-{index}"
            blocks.append(
                MaterialBlock(
                    priority=3,
                    path=logical,
                    block_type="image",
                    image_data_url=f"data:image/png;base64,{base64.b64encode(image_bytes).decode('ascii')}",
                    estimated_tokens=settings.DEFAULT_ESTIMATED_IMAGE_TOKENS,
                    logical_path=logical,
                    mime_hint="image/png",
                    origin="attachment",
                )
            )
    finally:
        document.close()
    return blocks


def _extract_ipynb_blocks(content: bytes, path: str) -> list[MaterialBlock]:
    try:
        notebook = json.loads(content.decode("utf-8"))
    except Exception:
        return []

    blocks: list[MaterialBlock] = []
    text_fragments: list[str] = []
    for index, cell in enumerate(notebook.get("cells") or [], start=1):
        cell_type = cell.get("cell_type") or "cell"
        source = "".join(cell.get("source") or [])
        if source.strip():
            text_fragments.append(f"## Cell {index} ({cell_type})\n{source.strip()}")
        for output in cell.get("outputs") or []:
            if output.get("output_type") == "stream":
                text = "".join(output.get("text") or [])
                if text.strip():
                    text_fragments.append(f"Output:\n{text.strip()[:MAX_IPYNB_OUTPUT_CHARS]}")
            data = output.get("data") or {}
            text_plain = data.get("text/plain")
            if text_plain:
                text = "".join(text_plain if isinstance(text_plain, list) else [str(text_plain)])
                if text.strip():
                    text_fragments.append(f"Output:\n{text.strip()[:MAX_IPYNB_OUTPUT_CHARS]}")
            image_png = data.get("image/png")
            if image_png:
                if isinstance(image_png, list):
                    image_png = "".join(image_png)
                logical_img = f"{path}#cell-{index}-output"
                blocks.append(
                    MaterialBlock(
                        priority=3,
                        path=logical_img,
                        block_type="image",
                        image_data_url=f"data:image/png;base64,{image_png}",
                        estimated_tokens=settings.DEFAULT_ESTIMATED_IMAGE_TOKENS,
                        logical_path=logical_img,
                        mime_hint="image/png",
                        origin="attachment",
                    )
                )
    if text_fragments:
        text, truncated = _truncate_text("\n\n".join(text_fragments))
        suffix = "\n\n[说明] Ipynb 文本输出已截断。" if truncated else ""
        mime_nb = "application/x-ipynb+json"
        blocks.insert(
            0,
            MaterialBlock(
                priority=2,
                path=path,
                block_type="text",
                text=f"### 附件（Jupyter 解析）\n**文件**: {path}\n**类型**: {mime_nb}\n\n{text}{suffix}",
                estimated_tokens=int(len(text) / 4) + 50,
                logical_path=path,
                mime_hint=mime_nb,
                origin="attachment",
                truncated=truncated,
            ),
        )
    return blocks


def _classify_and_extract(path: str, content: bytes) -> list[MaterialBlock]:
    suffix = PurePosixPath(path).suffix.lower()
    if suffix in PDF_EXTENSIONS:
        return _extract_pdf_images(content, path)
    if suffix in IPYNB_EXTENSIONS:
        return _extract_ipynb_blocks(content, path)
    if suffix in SUPPORTED_IMAGE_EXTENSIONS:
        mime_img = _guess_mime_type(path)
        return [
            MaterialBlock(
                priority=3,
                path=path,
                block_type="image",
                image_data_url=_bytes_to_data_url(path, content),
                estimated_tokens=settings.DEFAULT_ESTIMATED_IMAGE_TOKENS,
                logical_path=path,
                mime_hint=mime_img,
                origin="attachment",
            )
        ]
    if suffix == ".docx":
        text = _extract_docx_text(content)
    elif suffix in EXCEL_XLSX_EXTENSIONS:
        text = _extract_xlsx_text(content, path)
    elif suffix in EXCEL_XLS_EXTENSIONS:
        text = _extract_xls_text(content, path)
    elif suffix in LEGACY_WORD_DOC_EXTENSIONS:
        text = _extract_legacy_doc_text(content)
    elif suffix in SUPPORTED_TEXT_EXTENSIONS:
        text = _decode_bytes_as_text(content)
    else:
        return []
    text = text.strip()
    if not text:
        return []
    text, truncated = _truncate_text(text)
    suffix_note = "\n\n[说明] 文件内容过长，已截断。" if truncated else ""
    mime_doc = _guess_mime_type(path)
    return [
        MaterialBlock(
            priority=2,
            path=path,
            block_type="text",
            text=f"### 附件（解析文本）\n**文件**: {path}\n**类型**: {mime_doc}\n\n{text}{suffix_note}",
            estimated_tokens=int(len(text) / 4) + 50,
            logical_path=path,
            mime_hint=mime_doc,
            origin="attachment",
            truncated=truncated,
        )
    ]


def _walk_zip_bytes(
    content: bytes,
    *,
    root_path: str,
    depth: int,
    state: dict[str, Any],
) -> list[MaterialBlock]:
    blocks: list[MaterialBlock] = []
    if depth > MAX_ZIP_DEPTH:
        state["skipped"].append({"path": root_path, "reason": "超过最大嵌套深度"})
        return blocks
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            infos = sorted(archive.infolist(), key=lambda item: (item.filename or "").lower())
            for info in infos:
                if info.is_dir():
                    continue
                safe_child_path = _safe_relative_path(info.filename or "")
                if not safe_child_path:
                    state["skipped"].append({"path": f"{root_path}/{info.filename}", "reason": "非法路径"})
                    continue
                state["file_count"] += 1
                state["total_bytes"] += max(0, int(info.file_size or 0))
                if state["file_count"] > MAX_ZIP_FILES or state["total_bytes"] > MAX_ZIP_TOTAL_BYTES:
                    state["skipped"].append({"path": f"{root_path}/{safe_child_path}", "reason": "超出展开文件数或总大小限制"})
                    continue
                child_bytes = archive.read(info)
                child_path = f"{root_path}/{safe_child_path}"
                child_suffix = PurePosixPath(safe_child_path).suffix.lower()
                if child_suffix in ZIP_EXTENSIONS:
                    blocks.extend(_walk_zip_bytes(child_bytes, root_path=child_path, depth=depth + 1, state=state))
                elif child_suffix in RAR_EXTENSIONS:
                    blocks.extend(_walk_rar_bytes(child_bytes, root_path=child_path, depth=depth + 1, state=state))
                else:
                    child_blocks = _classify_and_extract(child_path, child_bytes)
                    if child_blocks:
                        blocks.extend(child_blocks)
                    else:
                        state["skipped"].append({"path": child_path, "reason": "无法识别或提取为空"})
    except zipfile.BadZipFile:
        state["skipped"].append({"path": root_path, "reason": "压缩包损坏或格式不合法"})
    return blocks


def _walk_rar_bytes(
    content: bytes,
    *,
    root_path: str,
    depth: int,
    state: dict[str, Any],
) -> list[MaterialBlock]:
    """
    RAR listing via rarfile; extraction via unrar/unrar-free or libarchive-backed tar.
    """
    blocks: list[MaterialBlock] = []
    if depth > MAX_ZIP_DEPTH:
        state["skipped"].append({"path": root_path, "reason": "超过最大嵌套深度"})
        return blocks
    if not _rar_extractor_tool_path():
        state["skipped"].append({"path": root_path, "reason": "RAR 解压需要安装 unrar、unrar-free 或 tar"})
        return blocks
    tmp_path: Optional[str] = None
    try:
        fd, tmp_path = tempfile.mkstemp(suffix=".rar")
        os.write(fd, content)
        os.close(fd)
        with rarfile.RarFile(tmp_path) as archive:
            if archive.needs_password():
                state["skipped"].append({"path": root_path, "reason": "RAR 已加密，不支持解压评分"})
                return blocks
            infos = sorted(archive.infolist(), key=lambda item: (item.filename or "").lower())
            for info in infos:
                if info.isdir():
                    continue
                if getattr(info, "needs_password", lambda: False)():
                    state["skipped"].append({"path": root_path, "reason": "RAR 内含加密条目，不支持"})
                    return blocks
                member_name = info.filename or ""
                safe_child_path = _safe_relative_path(member_name)
                if not safe_child_path:
                    state["skipped"].append({"path": f"{root_path}/{member_name}", "reason": "非法路径"})
                    continue
                state["file_count"] += 1
                state["total_bytes"] += max(0, int(info.file_size or 0))
                if state["file_count"] > MAX_ZIP_FILES or state["total_bytes"] > MAX_ZIP_TOTAL_BYTES:
                    state["skipped"].append({"path": f"{root_path}/{safe_child_path}", "reason": "超出展开文件数或总大小限制"})
                    continue
                child_path = f"{root_path}/{safe_child_path}"
                try:
                    child_bytes = _rar_read_member_bytes(tmp_path, member_name)
                except RuntimeError as exc:
                    state["skipped"].append({"path": child_path, "reason": str(exc)[:400]})
                    return blocks
                child_suffix = PurePosixPath(safe_child_path).suffix.lower()
                if child_suffix in ZIP_EXTENSIONS:
                    blocks.extend(_walk_zip_bytes(child_bytes, root_path=child_path, depth=depth + 1, state=state))
                elif child_suffix in RAR_EXTENSIONS:
                    blocks.extend(_walk_rar_bytes(child_bytes, root_path=child_path, depth=depth + 1, state=state))
                else:
                    child_blocks = _classify_and_extract(child_path, child_bytes)
                    if child_blocks:
                        blocks.extend(child_blocks)
                    else:
                        state["skipped"].append({"path": child_path, "reason": "无法识别或提取为空"})
    except rarfile.BadRarFile:
        state["skipped"].append({"path": root_path, "reason": "RAR 损坏或格式不合法"})
    except rarfile.NotRarFile:
        state["skipped"].append({"path": root_path, "reason": "不是有效的 RAR 文件"})
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    return blocks


def _collect_attachment_blocks(summary_path: str, attachment_name: str) -> tuple[list[MaterialBlock], list[dict[str, str]]]:
    file_path = get_attachment_file_path(summary_path)
    if not file_path or not file_path.exists():
        return [], [{"path": attachment_name or "attachment", "reason": "找不到原始附件文件"}]
    content = file_path.read_bytes()
    suffix = file_path.suffix.lower()
    state = {"file_count": 0, "total_bytes": 0, "skipped": []}
    name_lower = (attachment_name or "").lower()
    if suffix in ZIP_EXTENSIONS or name_lower.endswith(".zip"):
        blocks = _walk_zip_bytes(content, root_path=attachment_name or file_path.name, depth=1, state=state)
        return blocks, state["skipped"]
    if suffix in RAR_EXTENSIONS or name_lower.endswith(".rar"):
        blocks = _walk_rar_bytes(content, root_path=attachment_name or file_path.name, depth=1, state=state)
        return blocks, state["skipped"]
    blocks = _classify_and_extract(attachment_name or file_path.name, content)
    if blocks:
        return blocks, []
    return [], [{"path": attachment_name or file_path.name, "reason": "无法识别或提取为空"}]
