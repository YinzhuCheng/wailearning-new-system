"""Tests for LLM grading attachment parsing: Office formats and RAR."""

from __future__ import annotations

import io
import shutil
import zipfile
from pathlib import Path

import pytest
from docx import Document

from apps.backend.courseeval_backend.llm_grading import (
    _classify_and_extract,
    _walk_rar_bytes,
    _walk_zip_bytes,
)
from apps.backend.courseeval_backend.domains.llm.attachments import _rar_extractor_tool_path

_FIXTURE_RAR_DIR = Path(__file__).resolve().parents[2] / "fixtures" / "llm_rar"


def _require_rar_extractor() -> None:
    """RAR walking matches production: needs unrar/unrar-free or libarchive tar in PATH."""
    if not _rar_extractor_tool_path():
        pytest.skip("RAR extractor required in PATH (install unrar/unrar-free or provide libarchive tar)")


def _blocks_to_text(blocks: list) -> str:
    return "\n".join((b.text or "") for b in blocks if getattr(b, "block_type", None) == "text")


def test_extract_docx_includes_table_cells():
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Intro line for grading.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    doc.save(buf)
    blocks = _classify_and_extract("hw.docx", buf.getvalue())
    text = _blocks_to_text(blocks)
    assert "Intro line" in text
    assert "A1" in text and "B2" in text
    assert "[表格]" in text


def test_extract_xlsx_sheet_rows():
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "S1"
    ws.append(["Name", "Score"])
    ws.append(["Ann", 91])
    bio = io.BytesIO()
    wb.save(bio)
    blocks = _classify_and_extract("data.xlsx", bio.getvalue())
    text = _blocks_to_text(blocks)
    assert "工作表: S1" in text
    assert "Ann" in text and "91" in text


def test_extract_xls_fixture():
    from pathlib import Path

    p = Path(__file__).resolve().parents[2] / "fixtures" / "tiny.xls"
    assert p.is_file(), "fixtures/tiny.xls required (generated via xlwt)"
    blocks = _classify_and_extract("legacy.xls", p.read_bytes())
    text = _blocks_to_text(blocks)
    assert "HelloXls" in text
    assert "42" in text


def test_legacy_doc_non_ole_yields_empty():
    blocks = _classify_and_extract("fake.doc", b"not an ole file header")
    assert blocks == []


def test_zip_nested_xlsx():
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Z", 1])
    inner = io.BytesIO()
    wb.save(inner)
    zbuf = io.BytesIO()
    with zipfile.ZipFile(zbuf, "w") as zf:
        zf.writestr("inner/table.xlsx", inner.getvalue())
    state: dict = {"file_count": 0, "total_bytes": 0, "skipped": []}
    blocks = _walk_zip_bytes(zbuf.getvalue(), root_path="bundle.zip", depth=1, state=state)
    text = _blocks_to_text(blocks)
    assert "Z" in text


def test_rar_unencrypted_extracts_inner_txt():
    _require_rar_extractor()
    fixture = _FIXTURE_RAR_DIR / "unencrypted_nested_zip.rar"
    assert fixture.is_file(), f"missing committed fixture: {fixture}"
    content = fixture.read_bytes()
    state: dict = {"file_count": 0, "total_bytes": 0, "skipped": []}
    blocks = _walk_rar_bytes(content, root_path="t.rar", depth=1, state=state)
    text = _blocks_to_text(blocks)
    assert "RAR_INNER_UNIQUE_TEXT_12345" in text


def test_rar_password_rejected():
    _require_rar_extractor()
    fixture = _FIXTURE_RAR_DIR / "password_protected.rar"
    assert fixture.is_file(), f"missing committed fixture: {fixture}"
    content = fixture.read_bytes()
    state: dict = {"file_count": 0, "total_bytes": 0, "skipped": []}
    blocks = _walk_rar_bytes(content, root_path="enc.rar", depth=1, state=state)
    assert not blocks
    assert any("加密" in (s.get("reason") or "") for s in state["skipped"])
